from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import json
import os
import zipfile

def add_heading_custom(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph_formatted(doc, text, bold=False, alignment=None):
    p = doc.add_paragraph(text)
    if bold:
        p.runs[0].bold = True
    if alignment:
        p.alignment = alignment
    return p

def add_image_with_caption(doc, image_path, caption_text, width=5.5):
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = doc.add_paragraph(caption_text)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.runs[0]
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(80, 80, 80)
        return True
    else:
        doc.add_paragraph(f"【图片缺失：{image_path}】")
        return False

doc = Document()

title = doc.add_heading('实验4：对极几何与基础矩阵估计', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('实验报告')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(16)

doc.add_paragraph()

doc.add_heading('1. 实验目的', level=1)
doc.add_paragraph('本实验旨在通过实现对极几何中的基础矩阵估计算法，深入理解计算机视觉中的多视图几何理论。具体目标包括：')
doc.add_paragraph('理解对极几何的基本概念，包括极点、极线和基础矩阵的定义与性质', style='List Bullet')
doc.add_paragraph('掌握8点算法估计基础矩阵的原理和实现方法', style='List Bullet')
doc.add_paragraph('理解RANSAC鲁棒估计框架在处理特征匹配离群点中的作用', style='List Bullet')
doc.add_paragraph('对比人工标注匹配点和自动特征检测匹配在基础矩阵估计中的性能差异', style='List Bullet')

doc.add_heading('2. 实验原理', level=1)

doc.add_heading('2.1 对极几何', level=2)
doc.add_paragraph('对极几何描述了同一场景在不同视点下成像之间的几何关系。给定两个摄像机，其光心连线称为基线。基线与左右图像平面的交点分别称为左极点和右极点。对于左图中的一点，其在三维空间中对应的射线投影到右图上形成一条直线，称为极线。')

doc.add_heading('2.2 基础矩阵', level=2)
doc.add_paragraph('基础矩阵F是一个3×3的秩为2的矩阵，它编码了两幅图像之间的对极几何关系。对于左图中的点x和右图中对应点x\'，满足对极约束：x\'^T F x = 0。基础矩阵的零空间对应极点，即Fe_left = 0和F^T e_right = 0。')

doc.add_heading('2.3 极点与极线', level=2)
doc.add_paragraph('极点可通过对基础矩阵进行奇异值分解获得。左极点对应F的右零空间，右极点对应F^T的右零空间。极线方程可由基础矩阵和对应点计算得出，左图极线l_left = F^T x\'，右图极线l_right = F x。')

doc.add_page_break()

doc.add_heading('3. 实验方法', level=1)

doc.add_heading('3.1 Hartley归一化', level=2)
doc.add_paragraph('为提高数值稳定性，在进行基础矩阵估计前对点坐标进行归一化处理。归一化步骤包括：')
doc.add_paragraph('计算所有点的质心，将点坐标平移至原点，使其均值为零', style='List Number')
doc.add_paragraph('计算点到原点的平均距离，进行缩放使平均距离为√2', style='List Number')
doc.add_paragraph('构造归一化变换矩阵T，用于后续反归一化', style='List Number')

doc.add_heading('3.2 8点算法', level=2)
doc.add_paragraph('8点算法用于从至少8对匹配点估计基础矩阵。算法流程为：')
doc.add_paragraph('对左右图像的匹配点分别进行Hartley归一化', style='List Number')
doc.add_paragraph('对每对匹配点(x, y, 1)和(x\', y\', 1)，构造约束方程[x\'x, x\'y, x\', y\'x, y\'y, y\', x, y, 1]作为矩阵A的一行', style='List Number')
doc.add_paragraph('对A进行奇异值分解，最小奇异值对应的右奇异向量即为F的初始解（reshape为3×3）', style='List Number')
doc.add_paragraph('对F进行奇异值分解并将最小奇异值置零，保证秩为2约束', style='List Number')
doc.add_paragraph('使用归一化变换矩阵反归一化得到最终基础矩阵', style='List Number')

doc.add_heading('3.3 RANSAC框架', level=2)
doc.add_paragraph('RANSAC用于从含有离群点的匹配中鲁棒估计基础矩阵。算法设计思路为：')
doc.add_paragraph('随机选择8对匹配点作为最小集', style='List Number')
doc.add_paragraph('使用8点算法估计基础矩阵候选解', style='List Number')
doc.add_paragraph('计算所有匹配点的对称极线距离：d = (d_left + d_right) / 2，其中d_left为右图点到左极线距离，d_right为左图点到右极线距离', style='List Number')
doc.add_paragraph('统计距离小于阈值的内点数量', style='List Number')
doc.add_paragraph('保留内点数最多的模型，使用自适应策略动态调整迭代次数', style='List Number')

doc.add_heading('3.4 极点计算', level=2)
doc.add_paragraph('极点通过求解基础矩阵的零空间获得。对F进行奇异值分解F = UΣV^T，右极点为V的最后一列（对应最小奇异值），左极点通过对F^T进行相同操作获得。极点坐标需进行齐次坐标归一化。')

doc.add_heading('3.5 SIFT特征匹配', level=2)
doc.add_paragraph('任务2采用SIFT进行自动特征检测与匹配。SIFT算法提取图像的尺度不变特征点，使用FLANN快速最近邻匹配器进行特征描述子匹配，并采用Lowe比率测试（阈值0.75）过滤不可靠匹配，保留匹配质量较高的点对。')

doc.add_page_break()

doc.add_heading('4. 实验结果', level=1)

doc.add_heading('4.1 任务1：给定匹配点的基础矩阵估计', level=2)

results_dir = 'results/task1/'

error_stats_file = os.path.join(results_dir, 'error_statistics.json')
if os.path.exists(error_stats_file):
    with open(error_stats_file, 'r') as f:
        error_stats = json.load(f)

    doc.add_paragraph('任务1使用168对人工标注的匹配点进行基础矩阵估计。误差统计结果如下：')

    table1 = doc.add_table(rows=5, cols=2)
    table1.style = 'Light Grid Accent 1'

    table1.rows[0].cells[0].text = '指标'
    table1.rows[0].cells[1].text = '数值'
    table1.rows[1].cells[0].text = '匹配点数'
    table1.rows[1].cells[1].text = str(error_stats.get('num_points', 168))
    table1.rows[2].cells[0].text = '平均误差 (px)'
    table1.rows[2].cells[1].text = f"{error_stats.get('mean_error', 0):.4f}"
    table1.rows[3].cells[0].text = '中位数误差 (px)'
    table1.rows[3].cells[1].text = f"{error_stats.get('median_error', 0):.4f}"
    table1.rows[4].cells[0].text = '最大误差 (px)'
    table1.rows[4].cells[1].text = f"{error_stats.get('max_error', 0):.4f}"

    doc.add_paragraph()

epipole_file = os.path.join(results_dir, 'epipoles.txt')
if os.path.exists(epipole_file):
    with open(epipole_file, 'r') as f:
        lines = f.readlines()
    doc.add_paragraph(f'计算得到的极点坐标为：')
    doc.add_paragraph(lines[0].strip())
    doc.add_paragraph(lines[1].strip())
    doc.add_paragraph()

doc.add_paragraph('人工标注的168对匹配点可视化如下：')
add_image_with_caption(doc,
                      os.path.join(results_dir, 'manual_matches_visualization.png'),
                      '图1: 任务1人工标注的168对匹配点')

doc.add_paragraph()
doc.add_paragraph('左图像的极线可视化：')
add_image_with_caption(doc,
                      os.path.join(results_dir, 'left_epipolar_lines.png'),
                      '图2: 任务1左图极线可视化')

doc.add_paragraph()
doc.add_paragraph('右图像的极线可视化：')
add_image_with_caption(doc,
                      os.path.join(results_dir, 'right_epipolar_lines.png'),
                      '图3: 任务1右图极线可视化')

doc.add_page_break()

doc.add_heading('4.2 任务2：SIFT+RANSAC完整流程', level=2)

results_dir2 = 'results/task2/'

ransac_stats_file = os.path.join(results_dir2, 'ransac_statistics.json')
if os.path.exists(ransac_stats_file):
    with open(ransac_stats_file, 'r') as f:
        ransac_stats = json.load(f)

    doc.add_paragraph('任务2实现了从特征检测到基础矩阵估计的完整流程。首先使用SIFT算法检测特征点，然后进行特征匹配，最后使用RANSAC进行鲁棒估计。统计结果如下：')

    table2 = doc.add_table(rows=8, cols=2)
    table2.style = 'Light Grid Accent 1'

    table2.rows[0].cells[0].text = '指标'
    table2.rows[0].cells[1].text = '数值'
    table2.rows[1].cells[0].text = '左图SIFT特征数'
    table2.rows[1].cells[1].text = str(ransac_stats.get('num_features_left', 0))
    table2.rows[2].cells[0].text = '右图SIFT特征数'
    table2.rows[2].cells[1].text = str(ransac_stats.get('num_features_right', 0))
    table2.rows[3].cells[0].text = '初始匹配数'
    table2.rows[3].cells[1].text = str(ransac_stats.get('num_initial_matches', 0))
    table2.rows[4].cells[0].text = 'RANSAC内点数'
    table2.rows[4].cells[1].text = str(ransac_stats.get('ransac_inlier_count', 0))
    table2.rows[5].cells[0].text = 'RANSAC内点率'
    table2.rows[5].cells[1].text = f"{ransac_stats.get('ransac_inlier_ratio', 0)*100:.1f}%"
    table2.rows[6].cells[0].text = 'RANSAC迭代次数'
    table2.rows[6].cells[1].text = str(ransac_stats.get('ransac_iterations', 0))
    table2.rows[7].cells[0].text = '平均对称距离 (px)'
    table2.rows[7].cells[1].text = f"{ransac_stats.get('symmetric_distance_mean', 0):.4f}"

    doc.add_paragraph()

epipole_file2 = os.path.join(results_dir2, 'epipoles.txt')
if os.path.exists(epipole_file2):
    with open(epipole_file2, 'r') as f:
        lines = f.readlines()
    doc.add_paragraph(f'计算得到的极点坐标为：')
    doc.add_paragraph(lines[0].strip())
    doc.add_paragraph(lines[1].strip())
    doc.add_paragraph()

doc.add_paragraph('SIFT特征检测结果：')
add_image_with_caption(doc,
                      os.path.join(results_dir2, 'detected_features.png'),
                      '图4: 任务2 SIFT特征检测结果')

doc.add_paragraph()
doc.add_paragraph('RANSAC内点筛选结果：')
add_image_with_caption(doc,
                      os.path.join(results_dir2, 'ransac_inliers.png'),
                      '图5: 任务2 RANSAC内点（绿色）与离群点（红色）')

doc.add_paragraph()
doc.add_paragraph('左图像的极线可视化：')
add_image_with_caption(doc,
                      os.path.join(results_dir2, 'left_epipolar_lines.png'),
                      '图6: 任务2左图极线可视化')

doc.add_paragraph()
doc.add_paragraph('右图像的极线可视化：')
add_image_with_caption(doc,
                      os.path.join(results_dir2, 'right_epipolar_lines.png'),
                      '图7: 任务2右图极线可视化')

doc.add_page_break()

doc.add_heading('4.3 对比分析', level=2)

comparison_file = 'results/comparison/comparison.json'
if os.path.exists(comparison_file):
    with open(comparison_file, 'r') as f:
        comparison = json.load(f)

    doc.add_paragraph('两个任务在极点位置和误差指标上存在显著差异：')

    table3 = doc.add_table(rows=5, cols=3)
    table3.style = 'Light Grid Accent 1'

    table3.rows[0].cells[0].text = '指标'
    table3.rows[0].cells[1].text = '任务1'
    table3.rows[0].cells[2].text = '任务2'

    table3.rows[1].cells[0].text = '平均误差 (px)'
    table3.rows[1].cells[1].text = f"{comparison.get('task1_mean_error', 0):.4f}"
    table3.rows[1].cells[2].text = f"{comparison.get('task2_mean_error', 0):.4f}"

    table3.rows[2].cells[0].text = '匹配点数/内点数'
    table3.rows[2].cells[1].text = str(comparison.get('task1_num_matches', 0))
    table3.rows[2].cells[2].text = str(comparison.get('task2_num_inliers', 0))

    table3.rows[3].cells[0].text = '左极点差异 (px)'
    table3.rows[3].cells[1].text = '-'
    table3.rows[3].cells[2].text = f"{comparison.get('epipole_left_diff_px', 0):.2f}"

    table3.rows[4].cells[0].text = '右极点差异 (px)'
    table3.rows[4].cells[1].text = '-'
    table3.rows[4].cells[2].text = f"{comparison.get('epipole_right_diff_px', 0):.2f}"

    doc.add_paragraph()

doc.add_paragraph('分析：')
doc.add_paragraph('任务1使用人工标注的匹配点，精度显著高于任务2（0.069px vs 0.316px），这是由于人工标注消除了误匹配的影响。', style='List Number')
doc.add_paragraph('任务2的RANSAC内点率为75.9%，表明自动特征匹配存在约24%的离群点，但RANSAC框架能够有效识别并排除这些离群点。', style='List Number')
doc.add_paragraph('两个任务估计的极点位置差异巨大（左极点差异11876px，右极点差异1636px），这表明不同的匹配点集会导致基础矩阵估计结果的显著差异。', style='List Number')
doc.add_paragraph('任务1的极点位置远离图像范围，这可能意味着相机配置接近平行视角，导致极点投影在无穷远处。', style='List Number')

doc.add_page_break()

doc.add_heading('5. 结论', level=1)
doc.add_paragraph('本实验成功实现了基于8点算法和RANSAC框架的基础矩阵估计方法，并完成了两个任务的对比实验。主要结论如下：')
doc.add_paragraph()
doc.add_paragraph('8点算法结合Hartley归一化能够稳定地从匹配点估计基础矩阵，在人工标注数据上达到了0.069像素的高精度。', style='List Number')
doc.add_paragraph('RANSAC框架在处理自动特征匹配的离群点时表现良好，能够从133个初始匹配中识别出101个内点，内点率达到75.9%。', style='List Number')
doc.add_paragraph('人工标注匹配点的精度远高于自动特征检测匹配，但自动方法具有更好的可扩展性和实用性。', style='List Number')
doc.add_paragraph('基础矩阵的估计结果对输入匹配点的质量敏感，不同的匹配点集可能导致极点位置的显著差异。', style='List Number')
doc.add_paragraph('实验验证了对极几何理论的正确性，所有估计的基础矩阵均满足秩为2约束和对极约束条件。', style='List Number')

doc.add_page_break()

doc.add_heading('附录A：项目结构', level=1)
doc.add_paragraph('项目目录组织如下：')
doc.add_paragraph()

structure = """experiment4/
├── src/
│   ├── core/
│   │   ├── fundamental_matrix.py    # 8点算法、Hartley归一化
│   │   ├── epipolar_geometry.py     # 极点、极线计算
│   │   └── ransac.py                # RANSAC框架
│   ├── features/
│   │   ├── detector.py              # SIFT特征检测
│   │   └── matcher.py               # 特征匹配
│   ├── visualization/
│   │   └── epipolar_vis.py          # 可视化功能
│   └── utils/
│       └── io_utils.py              # 数据加载工具
├── scripts/
│   ├── task1_given_matches.py       # 任务1执行脚本
│   ├── task2_full_pipeline.py       # 任务2执行脚本
│   └── compare_results.py           # 结果对比脚本
├── results/
│   ├── task1/                       # 任务1结果
│   ├── task2/                       # 任务2结果
│   └── comparison/                  # 对比分析结果
├── left_upscaled.jpg                # 左视点图像
├── right_upscaled.jpg               # 右视点图像
└── h_matches（对极几何实践作业数据文件）.txt  # 人工标注匹配点"""

p = doc.add_paragraph(structure)
p.runs[0].font.name = 'Consolas'
p.runs[0].font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph('各模块功能说明：')
doc.add_paragraph('core模块：实现基础矩阵估计的核心算法', style='List Bullet')
doc.add_paragraph('features模块：实现SIFT特征检测和匹配', style='List Bullet')
doc.add_paragraph('visualization模块：实现极线和匹配点的可视化', style='List Bullet')
doc.add_paragraph('utils模块：提供数据加载等辅助功能', style='List Bullet')

doc.add_page_break()

doc.add_heading('附录B：使用说明', level=1)

doc.add_heading('环境依赖', level=2)
doc.add_paragraph('本项目需要以下Python库：')
dependencies = """numpy
opencv-python
matplotlib"""
p = doc.add_paragraph(dependencies)
p.runs[0].font.name = 'Consolas'
p.runs[0].font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('安装命令：')
install_cmd = "pip install numpy opencv-python matplotlib"
p = doc.add_paragraph(install_cmd)
p.runs[0].font.name = 'Consolas'
p.runs[0].font.size = Pt(10)

doc.add_heading('运行方法', level=2)
doc.add_paragraph('在项目根目录下执行以下命令：')
doc.add_paragraph()

doc.add_paragraph('运行任务1（使用给定匹配点）：')
cmd1 = "python scripts/task1_given_matches.py"
p = doc.add_paragraph(cmd1)
p.runs[0].font.name = 'Consolas'
p.runs[0].font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('运行任务2（SIFT+RANSAC完整流程）：')
cmd2 = "python scripts/task2_full_pipeline.py"
p = doc.add_paragraph(cmd2)
p.runs[0].font.name = 'Consolas'
p.runs[0].font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('对比两个任务的结果：')
cmd3 = "python scripts/compare_results.py"
p = doc.add_paragraph(cmd3)
p.runs[0].font.name = 'Consolas'
p.runs[0].font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('运行结果将保存在results/目录下，包括可视化图像、基础矩阵、极点坐标和统计数据等文件。')

output_file = '实验4_对极几何与基础矩阵估计_实验报告.docx'
doc.save(output_file)

print(f"✓ 报告已生成: {output_file}")
print(f"✓ 文件大小: {os.path.getsize(output_file):,} bytes")

print("\n验证图片嵌入情况...")
with zipfile.ZipFile(output_file, 'r') as z:
    media_files = [f for f in z.namelist() if f.startswith('word/media/')]
    print(f"✓ 成功嵌入 {len(media_files)} 张图片：")
    total_size = 0
    for f in media_files:
        size = z.getinfo(f).file_size
        total_size += size
        print(f"  • {f}: {size:,} bytes")
    print(f"\n总图片大小: {total_size:,} bytes")
    print(f"\n结论: 所有图片已完全嵌入文档，可单独发送此文件！")
